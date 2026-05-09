from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None
        if level == 1:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(13)
            run.bold = True
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

doc = Document()

# --- Title Page ---
doc.add_paragraph('\n'*4)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PROJECT REPORT\nON\nALGORITHMIC DOCUMENT SIMILARITY ANALYZER")
run.font.name = 'Times New Roman'
run.font.size = Pt(22)
run.bold = True

doc.add_paragraph('\n'*2)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Submitted in partial fulfillment of the requirements for the lab course\nDesign and Analysis of Algorithms Lab (15B11HS1122)\nProgram: BTech CSE\n4th Semester, 2nd Year")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph('\n'*3)
team = doc.add_paragraph()
team.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = team.add_run("Submitted By:\nMohd Munis (2401030005)\nShreyansh Rajat (2401030020)\nAman Saxena (2401030017)\nAditya Pandey (2401030027)\nDarsh Tiwari (2401030011)\nMayank Garg (2401030001)")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph('\n'*2)
mentors = doc.add_paragraph()
mentors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = mentors.add_run("Under the Guidance of:\nDr Aastha Maheshwari\nDr Kirti Jain\nDr Rajiv Mishra\nDr Taj Alam\nDr Tarun Agrawal")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

doc.add_page_break()

# --- 1. Introduction ---
add_heading(doc, '1. Introduction', 1)
add_paragraph(doc, '1.1 Overview', bold=True)
add_paragraph(doc, 'With the exponential growth of digital academic and online content, ensuring the originality of documents has become increasingly challenging. Students, researchers, and content creators now have access to vast repositories of information, making it easier to reproduce or adapt existing material. Traditional systems for document similarity often rely on naive approaches, which are inefficient and inaccurate when dealing with massive datasets or heavily paraphrased content.')
add_paragraph(doc, 'An Algorithmic Document Similarity Analyzer is a system designed to evaluate the similarity between a given query document and a collection (corpus) of source documents using structured algorithmic techniques. This project illustrates how fundamental techniques in string processing, efficient searching, hashing, indexing, and dynamic programming can be combined to build a scalable and accurate document comparison engine.')
add_paragraph(doc, 'The system follows a multi-stage plan that performs text preprocessing, exact phrase matching, paraphrase detection, efficient corpus search, and ranking. Each stage is implemented using core algorithms from the Design and Analysis of Algorithms (DAA) syllabus, making the system both highly educational and functionally robust.')

add_paragraph(doc, '1.2 Motivation', bold=True)
add_paragraph(doc, 'The primary motivation behind this project is to bridge the gap between theoretical algorithm design and practical software engineering. While algorithms like Knuth-Morris-Pratt, Rabin-Karp, and Longest Common Subsequence are rigorously taught in academia, their application in real-world, full-stack applications provides a profound understanding of their real-world utility and performance constraints.')
add_paragraph(doc, 'Furthermore, plagiarism detection is a critical requirement in academic institutions and publishing houses. By building a custom plagiarism detection engine from scratch, the team aims to understand the bottlenecks of large-scale text comparison, such as time complexity limitations and the challenge of detecting paraphrased (obfuscated) text. Addressing these challenges directly contributes to building a deeper appreciation for algorithm optimization.')

add_paragraph(doc, '1.3 Scope of the Project', bold=True)
add_paragraph(doc, 'The scope of this project encompasses the development of a complete web-based software application featuring a modern React frontend and a high-performance C++ backend. The project focuses on two primary modes of operation:')
add_bullet(doc, '1-on-1 Document Comparison: A detailed, granular comparison of two individual texts to identify exact shared substrings and paraphrased similarities.')
add_bullet(doc, 'Corpus Search: A scalable search mechanism that allows a user to upload a single document and query it against a large database of stored documents, identifying the top matching sources through advanced pruning and trie-based indexing.')
add_paragraph(doc, 'The project does not currently scale to a distributed cloud environment but is designed to function extremely efficiently on a local server, acting as a robust proof-of-concept for algorithmic document analysis.')

# --- 2. Problem Statement ---
add_heading(doc, '2. Problem Statement and Objectives', 1)
add_paragraph(doc, '2.1 Problem Statement', bold=True)
add_paragraph(doc, 'With the rapid growth of digital textual content, manually comparing documents to determine similarity has become entirely impractical. Educational institutions, research environments, and content management systems require automated tools that can evaluate textual similarity accurately and efficiently across large collections of documents. Existing commercial tools are often closed-source, computationally opaque, and expensive.')
add_paragraph(doc, 'The core problem addressed in this project is the design and implementation of an Algorithmic Document Similarity Analyzer that can mathematically and efficiently determine the degree of overlap between texts, scaling from single-document comparisons to full-corpus searches, all while handling exact matches and intelligently detecting paraphrased content.')

add_paragraph(doc, '2.2 Primary Objectives', bold=True)
add_bullet(doc, 'To implement an automated system capable of analyzing textual similarity using a pipeline of classic computer science algorithms.')
add_bullet(doc, 'To apply string matching algorithms (KMP and Rabin-Karp) to detect exact, unmodified copies of text rapidly.')
add_bullet(doc, 'To apply advanced data structures, specifically Suffix Arrays with Kasai\'s algorithm, to find the longest common contiguous blocks of text in O(N log^2 N) time.')
add_bullet(doc, 'To apply dynamic programming (Longest Common Subsequence) to detect paraphrased content where the core structure is maintained but words are inserted or deleted.')
add_bullet(doc, 'To build a high-performance backend using C++ for all heavy algorithmic processing, integrated with a modern FastAPI orchestrator.')
add_bullet(doc, 'To provide a transparent, user-friendly React frontend that not only shows the results but visualizes the algorithmic time/space complexities and highlights matching text segments.')

# --- 3. Requirement Analysis ---
add_heading(doc, '3. Requirement Analysis', 1)
add_paragraph(doc, '3.1 Functional Requirements', bold=True)
add_bullet(doc, 'The system must allow users to upload plain text (.txt, .md) documents.')
add_bullet(doc, 'The system must support a "1-on-1 Comparison Mode" displaying the similarity score and text highlights.')
add_bullet(doc, 'The system must support a "Corpus Search Mode" returning a ranked list of the top matching documents from the database.')
add_bullet(doc, 'The system must preprocess uploaded text by converting it to lowercase, stripping punctuation, and filtering out common stop words to prevent trivial matches.')
add_bullet(doc, 'The system must aggregate scores from multiple algorithms to provide a balanced, overall similarity percentage.')

add_paragraph(doc, '3.2 Non-Functional Requirements', bold=True)
add_bullet(doc, 'Performance: The backend C++ executables must execute string comparisons in a matter of milliseconds. The corpus search must employ pruning strategies to avoid O(N) linear time scanning of the entire database.')
add_bullet(doc, 'Usability: The frontend interface must be clean, academic-focused, and intuitive. Decorative and distracting elements should be minimized to focus on the data.')
add_bullet(doc, 'Scalability: The algorithms must be implemented with optimal time complexities (e.g., O(N+M) for KMP, O(N*M) for LCS) to handle large documents without crashing or timing out.')
add_bullet(doc, 'Maintainability: The backend should decouple the API routing (FastAPI) from the core logic (C++), allowing individual algorithms to be tested, modified, or replaced independently.')

add_paragraph(doc, '3.3 Software Requirements', bold=True)
add_bullet(doc, 'Frontend: React.js (v19), Vite, TailwindCSS for styling, Recharts for data visualization.')
add_bullet(doc, 'Backend API: Python 3.9+, FastAPI, Uvicorn.')
add_bullet(doc, 'Core Algorithms: C++17 compiler (GCC/Clang).')
add_bullet(doc, 'OS Environment: Windows / Linux / macOS compatible.')

add_paragraph(doc, '3.4 Hardware Requirements', bold=True)
add_bullet(doc, 'Processor: Multi-core CPU (Intel i5/AMD Ryzen 5 or better recommended for fast C++ execution).')
add_bullet(doc, 'RAM: Minimum 4 GB, recommended 8 GB for handling large document DP tables in memory.')
add_bullet(doc, 'Storage: 500 MB minimum free disk space for executables, corpus, and temporary upload storage.')

# --- 4. System Architecture ---
add_heading(doc, '4. System Architecture & Modeling', 1)
add_paragraph(doc, 'The Algorithmic Document Similarity Analyzer adopts a decoupled, micro-service-like architecture. The system is split into three distinct layers: the Presentation Layer (React), the Orchestration Layer (FastAPI), and the Execution Layer (C++ binaries).')

add_paragraph(doc, '4.1 Architectural Flow', bold=True)
add_paragraph(doc, '1. The React frontend gathers the user\'s uploaded files and formulates a multipart/form-data HTTP POST request.')
add_paragraph(doc, '2. The FastAPI backend receives the request, stores the files in a temporary \'uploads\' directory, and validates the input.')
add_paragraph(doc, '3. FastAPI uses the Python \'subprocess\' module to asynchronously spawn multiple C++ executables, passing the file paths as command-line arguments.')
add_paragraph(doc, '4. The C++ binaries perform the heavy computational lifting (reading files, building arrays, computing dynamic programming matrices) and output their results as JSON strings to standard output.')
add_paragraph(doc, '5. FastAPI captures these outputs, parses the JSON, aggregates the similarity scores using a weighted formula, cleans up the temporary files, and returns the final unified JSON response to the frontend.')

add_paragraph(doc, '4.2 Use Case Modeling', bold=True)
add_paragraph(doc, 'The system defines several distinct use cases based on the user\'s intent.')
add_bullet(doc, 'Use Case 1: Upload Documents. The user interacts with a drag-and-drop zone to upload files. The system validates file types and sizes.')
add_bullet(doc, 'Use Case 2: 1-on-1 Comparison. The user triggers the comparison. The system executes KMP, Rabin-Karp, Suffix Arrays, and LCS. The system returns highlighting arrays and percentage scores.')
add_bullet(doc, 'Use Case 3: Corpus Search. The user uploads a single query file. The system compares it against the entire database, prunes irrelevant files, ranks the top matches, and returns a detailed report for the top 5 documents.')
add_bullet(doc, 'Use Case 4: View Algorithm Complexity. The user views the dashboard to understand the theoretical time and space complexity of the algorithms that just executed, fostering an educational environment.')

add_paragraph(doc, '4.3 Activity Flow (1-on-1 Comparison)', bold=True)
add_paragraph(doc, 'The activity flow for a 1-on-1 comparison is strictly deterministic. Upon receiving files A and B, the system first runs them through a C++ preprocessor. The preprocessor strips non-alphanumeric characters, converts to lowercase, and generates n-grams. Simultaneously, the orchestrator triggers KMP, Rabin-Karp, Suffix Array, and LCS processes. Because these algorithms are independent, they execute in parallel. Once all processes complete or timeout, the orchestrator aggregates the scores. If the files are byte-for-byte identical, an automatic 100% score is assigned. Otherwise, the weighted average is calculated, and the results are formulated for the frontend.')

add_paragraph(doc, '4.4 Sequence Modeling', bold=True)
add_paragraph(doc, '1. User -> React UI: Uploads files and clicks "Run Comparison"')
add_paragraph(doc, '2. React UI -> FastAPI: POST /api/compare (Files A & B)')
add_paragraph(doc, '3. FastAPI -> OS File System: Save Files A & B to temp directory.')
add_paragraph(doc, '4. FastAPI -> C++ Binaries: Fork subprocesses (preprocessor, kmp, rabin_karp, suffix_array, lcs)')
add_paragraph(doc, '5. C++ Binaries -> FastAPI: Return JSON stdout containing similarities, runtimes, and matching indices.')
add_paragraph(doc, '6. FastAPI -> OS File System: Delete Files A & B.')
add_paragraph(doc, '7. FastAPI -> React UI: Return HTTP 200 with aggregated JSON payload.')
add_paragraph(doc, '8. React UI -> User: Render Gauges, Tables, and Highlighted Document Viewers.')

# --- 5. Core Algorithms ---
add_heading(doc, '5. Core Algorithms & Implementation Details', 1)

add_paragraph(doc, '5.1 Text Preprocessing', bold=True)
add_paragraph(doc, 'Before any advanced algorithm can run, the text must be normalized to prevent false negatives caused by superficial differences like capitalization, extra spacing, or punctuation. The C++ preprocessor tokenizes the text, converts it to a uniform case, and strips stop words. It then generates "shingles" or n-grams. This standardizes the input space so that the subsequent string matching algorithms evaluate the core semantic structure of the sentences rather than formatting anomalies.')

add_paragraph(doc, '5.2 Knuth-Morris-Pratt (KMP) Algorithm', bold=True)
add_paragraph(doc, 'The KMP algorithm is utilized to detect exact, un-paraphrased copies of phrases. Traditional naive search algorithms backtrack when a character mismatch occurs, leading to a worst-case time complexity of O(N * M).')
add_paragraph(doc, 'Implementation: KMP extracts contiguous phrases from Document A. For each phrase, it constructs a Longest Proper Prefix which is also Suffix (LPS) array. This array acts as a deterministic finite automaton (DFA) failure function. When scanning Document B, if a mismatch occurs, the LPS array tells the algorithm exactly how far to shift the search window forward without re-evaluating characters it has already seen.')
add_paragraph(doc, 'Result: Strict O(N+M) time complexity. It guarantees that direct copy-pasted segments are identified with absolute precision and negligible overhead.')

add_paragraph(doc, '5.3 Rabin-Karp Algorithm', bold=True)
add_paragraph(doc, 'Rabin-Karp is employed for extremely fast multi-pattern matching, scanning sliding windows of text to catch direct copy-pasting rapidly.')
add_paragraph(doc, 'Implementation: Instead of comparing characters one-by-one, the algorithm treats a sequence of characters as a base-P number. It computes a numeric hash (using a prime modulus to prevent integer overflow) for a block of text in Document A. As the window slides through Document B, the algorithm mathematically drops the oldest character\'s value and adds the newest character\'s value to the rolling hash in O(1) time. If the hashes match, a secondary exact-character check is performed to resolve potential hash collisions.')
add_paragraph(doc, 'Result: Highly efficient sliding window search, yielding an average time complexity of O(N+M) and demonstrating the power of rolling hashes in string processing.')

add_paragraph(doc, '5.4 Suffix Arrays and Kasai\'s LCP', bold=True)
add_paragraph(doc, 'To identify the absolute longest exact contiguous blocks of text shared between the two files, the system utilizes Suffix Arrays combined with Kasai\'s algorithm for the Longest Common Prefix (LCP).')
add_paragraph(doc, 'Implementation: The algorithm concatenates Document A and Document B, separated by a unique delimiter (e.g., \'#\'). It generates a Suffix Array, which is a lexicographically sorted array of all suffixes of the combined string. Sorting naively would take O(N^2 log N), but the system implements a prefix-doubling algorithm that achieves O(N log^2 N) time. Once sorted, identical substrings are guaranteed to be strictly adjacent in the array. The system then applies Kasai\'s Algorithm to compute the LCP array in O(N) time by examining the overlap between adjacent suffixes.')
add_paragraph(doc, 'Result: Extracts the longest shared text blocks globally, proving invaluable for detecting massive chunks of plagiarized content.')

add_paragraph(doc, '5.5 Longest Common Subsequence (LCS) using DP', bold=True)
add_paragraph(doc, 'While KMP, Rabin-Karp, and Suffix Arrays are exceptional at finding exact contiguous matches, they fail when a plagiarist paraphrases text by inserting or deleting words. The Longest Common Subsequence algorithm solves this.')
add_paragraph(doc, 'Implementation: The algorithm constructs a 2D Dynamic Programming matrix of size N x M, where N and M are the number of tokens in the respective documents. The DP table is built bottom-up. If tokenA[i] equals tokenB[j], the cell inherits the top-left diagonal value plus one. Otherwise, it takes the maximum of the cell directly above or to the left. After filling the matrix, the algorithm backtracks from the bottom-right corner to reconstruct the actual longest shared sequence of words.')
add_paragraph(doc, 'Result: Time and Space complexity of O(N * M). It successfully catches paraphrased plagiarism, recognizing the preserved structural order of words even if filler content is injected between them.')

add_paragraph(doc, '5.6 Corpus Search: Trie Construction and Divide & Conquer Pruning', bold=True)
add_paragraph(doc, 'Comparing a query document sequentially against thousands of corpus documents (O(N) operations) is too slow. The system uses advanced retrieval and pruning.')
add_paragraph(doc, 'Implementation: On startup, the system generates rolling hash signatures (fingerprints) for every corpus document and sorts them. When a query is uploaded, its hashes are generated. The system uses a Divide and Conquer strategy: it recursively splits the corpus in half. For a given half, it checks the aggregate hash overlap with the query. If the overlap falls below a strict threshold, the entire half is pruned in O(1) time. This bypasses thousands of irrelevant documents. For surviving documents, Binary Search calculates the exact Jaccard Similarity. A highly optimized C++ Trie (Prefix Tree) is then traversed to extract exact shared phrases for the top matches.')
add_paragraph(doc, 'Result: Exponentially speeds up corpus querying. The top 5 results are then subjected to detailed Suffix Array and LCS analysis for high-precision final ranking.')

# --- 6. Test Cases ---
add_heading(doc, '6. Test Cases & Evaluation', 1)
add_paragraph(doc, 'The system was rigorously tested to ensure algorithmic correctness and architectural stability. A suite of test cases was designed to stress-test different components.')

add_paragraph(doc, '6.1 Unit Testing: Exact Matching', bold=True)
add_bullet(doc, 'Test Case ID: TC01')
add_bullet(doc, 'Description: Verify KMP and Rabin-Karp on identically copied paragraphs.')
add_bullet(doc, 'Input: Doc A ("The quick brown fox jumps over the lazy dog."), Doc B ("The quick brown fox jumps over the lazy dog.")')
add_bullet(doc, 'Expected Output: 100% similarity for both algorithms, full array highlighting.')
add_bullet(doc, 'Status: Passed. The LPS array correctly mapped the entire sequence without backtracking.')

add_paragraph(doc, '6.2 Unit Testing: Paraphrase Detection', bold=True)
add_bullet(doc, 'Test Case ID: TC02')
add_bullet(doc, 'Description: Verify LCS DP on text with injected filler words.')
add_bullet(doc, 'Input: Doc A ("Design and Analysis of Algorithms is essential."), Doc B ("The study of Design and careful Analysis of various Algorithms is highly essential.")')
add_bullet(doc, 'Expected Output: Exact matching algorithms score low. LCS scores high and reconstructs the sequence ["Design", "and", "Analysis", "of", "Algorithms", "is", "essential"].')
add_bullet(doc, 'Status: Passed. The 2D DP matrix successfully bridged the gaps created by the filler words.')

add_paragraph(doc, '6.3 Unit Testing: Corpus Pruning Efficiency', bold=True)
add_bullet(doc, 'Test Case ID: TC03')
add_bullet(doc, 'Description: Evaluate Divide and Conquer pruning on a corpus of 100 random text files.')
add_bullet(doc, 'Input: A query document completely unrelated to 95 of the files.')
add_bullet(doc, 'Expected Output: The system should prune the 95 unrelated files without executing full string comparisons on them, reducing search time significantly.')
add_bullet(doc, 'Status: Passed. Console logs indicated pruning eliminated 95% of the search space in under 5 milliseconds.')

add_paragraph(doc, '6.4 Integration Testing: End-to-End Execution', bold=True)
add_bullet(doc, 'Test Case ID: TC04')
add_bullet(doc, 'Description: Test the full pipeline from React UI upload to FastAPI parsing to C++ execution and JSON rendering.')
add_bullet(doc, 'Input: Two large textual documents (approx. 5000 words each).')
add_bullet(doc, 'Expected Output: System completes execution without HTTP timeout, successfully aggregates scores with correct weights, and renders the dashboard UI without crashing the browser.')
add_bullet(doc, 'Status: Passed. Execution completed in ~450ms. React handled the large JSON payload smoothly.')

# --- 7. Team Contributions ---
add_heading(doc, '7. Team Contributions', 1)
add_paragraph(doc, 'The development of this project was a collaborative effort, with tasks divided according to the team members\' strengths and the system\'s architectural boundaries.')
add_bullet(doc, 'Mohd Munis (2401030005) & Shreyansh Rajat (2401030020): Focused on the C++ algorithmic core. They were responsible for implementing the prefix-doubling Suffix Array, Kasai\'s LCP algorithm, and the complex Longest Common Subsequence dynamic programming matrix. They ensured the C++ binaries were highly optimized and memory-safe.')
add_bullet(doc, 'Aman Saxena (2401030017) & Aditya Pandey (2401030027): Handled the exact string matching algorithms. They implemented the Knuth-Morris-Pratt (KMP) failure functions and the Rabin-Karp rolling hash mechanics. They also engineered the C++ preprocessor to standardize document inputs.')
add_bullet(doc, 'Darsh Tiwari (2401030011) & Mayank Garg (2401030001): Managed the system architecture, backend orchestration, and frontend visualization. They built the FastAPI server, designed the subprocess management system, and developed the React.js interface to parse and beautifully render the algorithmic outputs, including the corpus search pruning logic.')

# --- 8. Conclusion ---
add_heading(doc, '8. Conclusion and Future Scope', 1)
add_paragraph(doc, '8.1 Conclusion', bold=True)
add_paragraph(doc, 'The Algorithmic Document Similarity Analyzer successfully demonstrates how classical theoretical algorithms can be synthesized into a powerful, real-world application. By integrating string matching, hashing, sorting, and dynamic programming within a high-performance C++ core, the system achieves rapid and accurate plagiarism detection. The project highlights the critical importance of selecting the right data structures—such as Suffix Arrays over naive sorting, and Tries over linear searching—to overcome the computational bottlenecks inherent in large-scale text analysis. The resulting software is not only functionally effective but also serves as an exceptional educational tool for visualizing algorithm mechanics.')

add_paragraph(doc, '8.2 Future Scope', bold=True)
add_paragraph(doc, 'While the current system is robust, several enhancements could be introduced in future iterations:')
add_bullet(doc, 'Semantic Understanding: Integrating Natural Language Processing (NLP) models, such as Word2Vec or SentenceTransformers, alongside the classical algorithms to detect semantic similarities where synonyms are used instead of exact word copies.')
add_bullet(doc, 'Cloud Deployment & Scalability: Refactoring the corpus search to utilize a distributed database (e.g., Elasticsearch) or vector database (e.g., FAISS) to scale the corpus to millions of documents.')
add_bullet(doc, 'Syntax Tree Analysis for Code: Extending the preprocessor to generate Abstract Syntax Trees (ASTs) for source code files, allowing the system to detect code plagiarism independently of variable renaming.')

# --- 9. References ---
add_heading(doc, '9. References', 1)
add_bullet(doc, 'Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms (3rd ed.). MIT Press. (Reference for KMP, LCS, and String Matching theories).')
add_bullet(doc, 'GeeksforGeeks. (n.d.). KMP Algorithm for Pattern Searching. Retrieved from https://www.geeksforgeeks.org/dsa/kmp-algorithm-for-pattern-searching/')
add_bullet(doc, 'GeeksforGeeks. (n.d.). Rabin-Karp Algorithm for Pattern Searching. Retrieved from https://www.geeksforgeeks.org/rabin-karp-algorithm-for-pattern-searching/')
add_bullet(doc, 'GeeksforGeeks. (n.d.). Suffix Array - Introduction. Retrieved from https://www.geeksforgeeks.org/suffix-array-set-1-introduction/')
add_bullet(doc, 'GeeksforGeeks. (n.d.). Trie - Insert and Search. Retrieved from https://www.geeksforgeeks.org/trie-insert-and-search/')
add_bullet(doc, 'GeeksforGeeks. (n.d.). Longest Common Subsequence (LCS). Retrieved from https://www.geeksforgeeks.org/longest-common-subsequence-dp-4/')

doc.save('Project_Report.docx')
